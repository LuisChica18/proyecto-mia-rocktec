"""
generar_informe_docx.py
Genera 06_resultados/INFORME_AJUSTES_Y_VALIDACION.docx — Rocktec MIA 2026

Reproduce el documento de entrega (ajustes de prototipo, reporte de análisis
estadístico, consolidación de documentación y síntesis reflexiva) a partir de:
  - 06_resultados/comparacion_modelos.json      (LR/SVM entrenados sobre consenso)
  - 06_resultados/validacion_estadistica.json   (CV pareada baseline vs. ajustado)
  - requirements.txt                            (dependencias)
  - git log                                     (evidencia de control de versiones)

Uso: python 02_scripts/generar_informe_docx.py   (ejecutar desde la raíz del repo)
"""

import json
import subprocess
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RAIZ = Path(__file__).parent.parent
RUTA_COMPARACION = RAIZ / '06_resultados' / 'comparacion_modelos.json'
RUTA_VALIDACION = RAIZ / '06_resultados' / 'validacion_estadistica.json'
RUTA_REQS = RAIZ / 'requirements.txt'
RUTA_SALIDA = RAIZ / '06_resultados' / 'INFORME_AJUSTES_Y_VALIDACION.docx'
RUTA_CM_LR = RAIZ / '06_resultados' / 'modelos' / 'confusion_matrix_logistic_regression.png'
RUTA_CM_SVM = RAIZ / '06_resultados' / 'modelos' / 'confusion_matrix_svm.png'
RUTA_BOXPLOT = RAIZ / '06_resultados' / 'modelos' / 'validacion_cv_boxplot.png'

AZUL = RGBColor(0x1F, 0x3A, 0x5F)


def git_log(n=12):
    out = subprocess.run(
        ['git', '-C', str(RAIZ), 'log', f'-{n}', '--pretty=format:%h|%ad|%s', '--date=short'],
        capture_output=True, text=True
    ).stdout.strip()
    return [line.split('|', 2) for line in out.splitlines() if line]


def heading(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


def tabla_simple(doc, encabezados, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = 'Light Grid Accent 1'
    for i, texto in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = texto
        celda.paragraphs[0].runs[0].bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = str(valor)
    return t


def main():
    comparacion = json.loads(RUTA_COMPARACION.read_text(encoding='utf-8'))
    validacion = json.loads(RUTA_VALIDACION.read_text(encoding='utf-8'))
    resultados_cv = validacion['resultados']
    comparaciones_cv = validacion['comparaciones']
    meta = validacion['metodologia']

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    # ── Portada ──
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('Ajustes al Prototipo, Validación Estadística\ny Documentación del Repositorio')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = AZUL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('Plataforma de Inteligencia Comercial para Rocktec — MIA 2026\n'
                 'Universidad de las Américas (UDLA)\n'
                 f'Fecha: {date.today().isoformat()}').italic = True

    doc.add_page_break()

    # ── 1. Ajustes de prototipo ──
    heading(doc, '1. Ajustes de Prototipo')

    doc.add_paragraph(
        'Al analizar los resultados preliminares se detectó una falla de causa raíz en el '
        'preprocesamiento: el script calcular_kappa.py calculaba el coeficiente de Cohen\'s Kappa '
        'entre los 3 anotadores, pero nunca escribía el archivo dataset_consenso_final.csv que el '
        'pipeline de feature engineering (04_feature_engineering.py) esperaba como fuente principal '
        'de etiquetas. Adicionalmente, la ruta de respaldo apuntaba a un nombre de archivo y una '
        'columna que no coincidían con los reales. Como consecuencia, todos los entrenamientos '
        'anteriores usaban etiquetas heurísticas (por palabras clave) en lugar de las 1,500 '
        'anotaciones humanas reales, a pesar de que el Kappa de 0.8851 ya certificaba su calidad.'
    )

    heading(doc, '1.1 Qué se ajustó', nivel=2)
    for texto in [
        'Preprocesamiento: se implementó la generación real del dataset de consenso (voto '
        'mayoritario 2/3 entre los 3 anotadores) en calcular_kappa.py, produciendo '
        '04_anotaciones/dataset_consenso_final.csv (1,297 de 1,500 filas con consenso válido; '
        '203 quedaron sin consenso por desacuerdo total o anotaciones inválidas). Sobre las mismas '
        'filas se calculó también una etiqueta heurística baseline, para permitir comparaciones '
        'pareadas exactas.',
        'Hiperparámetros: se ampliaron las grillas de búsqueda de GridSearchCV '
        '(Logistic Regression: C de 4 a 5 valores, 0.01–100; LinearSVC: C de 4 a 6 valores, '
        '0.001–100), y se hizo el número de folds de validación cruzada adaptativo al tamaño de '
        'la clase más pequeña del conjunto de entrenamiento (clases como QUE y SEG tienen muy '
        'pocos ejemplos en el dataset de consenso).',
        'Modelado: se ajustó el uso de CalibratedClassifierCV para LinearSVC de modo que solo se '
        'active cuando hay suficientes ejemplos por clase para calibrar probabilidades; en caso '
        'contrario se entrena LinearSVC sin calibrar, evitando fallas de entrenamiento en clases '
        'raras sin sacrificar la métrica de evaluación (F1-macro), que solo usa predicciones '
        'duras.',
        'Validación estadística: se implementó un nuevo script (07_validacion_estadistica.py) que '
        'aplica un diseño de validación cruzada pareado para comparar rigurosamente el modelo '
        'baseline (heurístico) contra el ajustado (consenso), descrito en la sección 2.',
    ]:
        doc.add_paragraph(texto, style='List Bullet')

    heading(doc, '1.2 Por qué eran necesarios', nivel=2)
    doc.add_paragraph(
        'Sin corregir el gap de preprocesamiento, todo el trabajo de anotación manual (Fase 1B, '
        '3 anotadores, ~4-6 horas cada uno) no tenía ningún efecto sobre el modelo entrenado: se '
        'estaba optimizando y reportando métricas de un problema distinto (predecir una regla de '
        'palabras clave) al que realmente importa (predecir el juicio humano de intención). Los '
        'ajustes de hiperparámetros y calibración fueron necesarios porque el dataset de consenso '
        'real es mucho más pequeño (1,297 filas) y desbalanceado que el dataset heurístico '
        '(~9,300 filas), lo que expone fallas de configuración que antes no se manifestaban.'
    )

    heading(doc, '1.3 Qué se esperaba mejorar', nivel=2)
    doc.add_paragraph(
        'Se esperaba obtener una estimación honesta y reproducible del desempeño del modelo sobre '
        'el criterio humano real (no sobre una regla heurística que se autovalida), junto con una '
        'cuantificación rigurosa de la incertidumbre de esa estimación mediante validación cruzada '
        'repetida, en lugar de depender de un único split train/test.'
    )

    # ── 2. Reporte de análisis estadístico ──
    doc.add_page_break()
    heading(doc, '2. Reporte de Análisis Estadístico')

    heading(doc, '2.1 Metodología', nivel=2)
    doc.add_paragraph(
        f'Se diseñó una validación cruzada PAREADA: RepeatedKFold con '
        f'{meta["n_splits"]} particiones y {meta["n_repeats"]} repeticiones '
        f'({meta["n_splits"] * meta["n_repeats"]} folds en total), sin estratificar, aplicada '
        f'sobre las MISMAS {meta["n_filas"]} filas anotadas. Al no depender la partición de la '
        'variable objetivo, los mismos índices de fila (mismo train/test) se reutilizan para '
        'evaluar tanto la etiqueta heurística (baseline) como la de consenso (ajustado), aislando '
        'así el efecto de la calidad del etiquetado del efecto del tamaño de la muestra — que sí '
        'difiere entre el dataset heurístico completo (~9,300 filas) y el de consenso (1,297). '
        'Los hiperparámetros de LR (C=0.1) y SVM (C=1.0) se fijaron en los valores ya seleccionados '
        'por GridSearchCV, y el vectorizador TF-IDF se reajustó en cada fold usando solo el texto '
        'de entrenamiento (sin fuga de información hacia el test). La métrica reportada es '
        'F1-macro sobre las 7 clases fijas del catálogo (zero_division=0), para que el '
        'denominador del promedio macro no cambie según qué clases aparecen en cada fold.'
    )

    heading(doc, '2.2 Comparación baseline (heurístico) vs. ajustado (consenso)', nivel=2)
    filas_cv = []
    etiquetas = {'lr_heuristico': 'LR — heurístico (baseline)',
                 'lr_consenso': 'LR — consenso (ajustado)',
                 'svm_heuristico': 'SVM — heurístico (baseline)',
                 'svm_consenso': 'SVM — consenso (ajustado)'}
    for clave, etiqueta in etiquetas.items():
        r = resultados_cv[clave]['f1_macro']
        filas_cv.append([
            etiqueta, f"{r['media']:.4f}", f"{r['std']:.4f}",
            f"[{r['ic95_low']:.4f}, {r['ic95_high']:.4f}]", r['n'],
        ])
    tabla_simple(doc, ['Modelo / fuente de datos', 'F1-macro (media)', 'Desv. estándar',
                        'IC 95%', 'n folds'], filas_cv)

    doc.add_paragraph()
    doc.add_paragraph(
        f'Como referencia adicional, un único split train/test (80/20) sobre el dataset completo '
        f'de consenso (1,297 filas) dio F1-macro = '
        f'{comparacion["logistic_regression"]["f1_macro"]:.4f} (LR) y '
        f'{comparacion["linear_svc"]["f1_macro"]:.4f} (SVM) — consistente en orden de magnitud '
        f'con las medias de la validación cruzada pareada de la tabla anterior.'
    )

    heading(doc, '2.3 Pruebas de hipótesis (comparaciones pareadas)', nivel=2)
    filas_test = []
    nombres_test = {
        'lr_consenso_vs_heuristico': 'LR: consenso vs. heurístico',
        'svm_consenso_vs_heuristico': 'SVM: consenso vs. heurístico',
        'svm_vs_lr_consenso': 'Consenso: SVM vs. LR',
    }
    for clave, nombre in nombres_test.items():
        c = comparaciones_cv[clave]
        filas_test.append([
            nombre, f"{c['diferencia_media']:+.4f}",
            f"{c['wilcoxon_statistic']:.1f}", f"{c['wilcoxon_pvalue']:.4f}",
            f"{c['ttest_pareado_statistic']:.3f}", f"{c['ttest_pareado_pvalue']:.4f}",
        ])
    tabla_simple(doc, ['Comparación', 'Δ media F1-macro', 'Wilcoxon (stat)', 'Wilcoxon (p)',
                        't-test pareado (stat)', 't-test pareado (p)'], filas_test)

    doc.add_paragraph()
    doc.add_paragraph(
        'Ambas pruebas (Wilcoxon signed-rank, no paramétrica, y t-test pareado como verificación '
        'cruzada) coinciden: las diferencias son estadísticamente significativas (p < 0.001 en '
        'los 3 casos). El modelo entrenado sobre etiquetas heurísticas obtiene un F1-macro '
        'significativamente mayor que el entrenado sobre el consenso humano — incluso controlando '
        'el tamaño de muestra — y, dentro del brazo de consenso, LR supera significativamente a '
        'SVM (lo opuesto de lo observado con las etiquetas heurísticas, donde SVM parecía muy '
        'superior).'
    )

    heading(doc, '2.4 Gráficos', nivel=2)
    if RUTA_BOXPLOT.exists():
        doc.add_picture(str(RUTA_BOXPLOT), width=Inches(5.5))
    if RUTA_CM_LR.exists():
        doc.add_picture(str(RUTA_CM_LR), width=Inches(4.5))
    if RUTA_CM_SVM.exists():
        doc.add_picture(str(RUTA_CM_SVM), width=Inches(4.5))

    heading(doc, '2.5 Interpretación de los hallazgos', nivel=2)
    doc.add_paragraph(
        'El hallazgo central es que el buen desempeño reportado originalmente (SVM F1-macro '
        '0.936 sobre etiquetas heurísticas) era en gran parte un artefacto de circularidad: las '
        'mismas reglas de palabras clave que generaban la etiqueta (p.ej. "cotizacion", '
        '"presupuesto" → COT) son triviales de recuperar para un modelo bag-of-words como TF-IDF + '
        'SVM/LR, por lo que la métrica medía qué tan bien el modelo reconstruía la regla, no qué '
        'tan bien predice la intención real del cliente. Al evaluar sobre el consenso humano —que '
        'incorpora matices, contexto y desacuerdo genuino entre anotadores (8.8% de desacuerdo '
        'total)— el desempeño honesto es sustancialmente menor (F1-macro ≈ 0.46–0.59) y ninguno '
        'de los dos modelos alcanza la meta de F1 ≥ 0.75. Esto no es un retroceso: es la primera '
        'medición fiable del punto de partida real para Fase 2+, y evita optimizar la arquitectura '
        'MLOps sobre una métrica engañosa.'
    )

    # ── 3. Consolidación y documentación del repositorio ──
    doc.add_page_break()
    heading(doc, '3. Consolidación y Documentación del Repositorio')

    heading(doc, '3.1 Cambios de documentación', nivel=2)
    for texto in [
        'README.md: institución corregida a Universidad de las Américas (UDLA); nueva sección de '
        'stack tecnológico; tabla de resultados baseline vs. ajustado; fila de Fase 1C en el '
        'cronograma; enlaces a este informe y a CHANGELOG.md.',
        'README_PATY.md y README_LuisC.md: eliminados por redundantes (contenido único fusionado '
        'en README.md; el segundo estaba vacío).',
        '03_datos_procesados/README.md: actualizado — describía solo el pipeline original de 2 '
        'fuentes (9,317 filas); ahora documenta también el pipeline de 4 bases y el estado real de '
        'Fase 1B/1C.',
        'CLAUDE.md: institución corregida a UDLA; documentado el comportamiento corregido de '
        'calcular_kappa.py y los nuevos scripts 07_validacion_estadistica.py y '
        'generar_informe_docx.py; documentadas las variables de entorno necesarias '
        '(MLFLOW_ALLOW_FILE_STORE) para este entorno de ejecución.',
        'CHANGELOG.md (nuevo): bitácora cronológica reconstruida desde el historial de git '
        '(Fase 1 → Fase 1B) más la entrada detallada de esta Fase 1C.',
    ]:
        doc.add_paragraph(texto, style='List Bullet')

    heading(doc, '3.2 Dependencias', nivel=2)
    if RUTA_REQS.exists():
        deps = [l.strip() for l in RUTA_REQS.read_text(encoding='utf-8').splitlines() if l.strip()]
        filas_deps = []
        for linea in deps:
            if '==' in linea:
                nombre, version = linea.split('==', 1)
            elif '>=' in linea:
                nombre, version = linea.split('>=', 1)
                version = f'>= {version}'
            else:
                nombre, version = linea, '—'
            filas_deps.append([nombre, version])
        tabla_simple(doc, ['Paquete', 'Versión'], filas_deps)

    heading(doc, '3.3 Instrucciones de ejecución (resumen)', nivel=2)
    for paso in [
        'pip install -r requirements.txt',
        'python 02_scripts/calcular_kappa.py "04_anotaciones/ROCKTEC_BASE_FINAL_ANOTACION_1500 1ERA ETIQUETA.xlsx"  '
        '→ genera dataset_consenso_final.csv',
        'python 02_scripts/05_entrenar_modelos.py --skip-beto  '
        '(usar MLFLOW_ALLOW_FILE_STORE=true si mlflow >= 2.11 bloquea el backend de archivos)',
        'python 02_scripts/07_validacion_estadistica.py  → validación cruzada pareada',
        'python 02_scripts/generar_informe_docx.py  → regenera este documento',
    ]:
        doc.add_paragraph(paso, style='List Number')

    heading(doc, '3.4 Evidencia de control de versiones (commits)', nivel=2)
    filas_git = [[h, f, m[:90]] for h, f, m in git_log(12)]
    tabla_simple(doc, ['Commit', 'Fecha', 'Mensaje'], filas_git)

    # ── 4. Síntesis reflexiva ──
    doc.add_page_break()
    heading(doc, '4. Síntesis Reflexiva')

    heading(doc, '4.1 Decisiones técnicas más relevantes', nivel=2)
    for texto in [
        'Diseño pareado (mismas filas, mismos folds) para comparar baseline vs. ajustado: sin '
        'este control, la comparación habría confundido dos efectos distintos — calidad de la '
        'etiqueta y tamaño de la muestra (9,300 vs. 1,297 filas) — invalidando cualquier '
        'conclusión causal sobre por qué cambió el desempeño.',
        'Prueba de Wilcoxon signed-rank como prueba principal (con t-test pareado como '
        'verificación cruzada): al no poder garantizar normalidad de las diferencias fold-a-fold '
        'con solo 25 observaciones, se prefirió una prueba no paramétrica más robusta.',
        'Alcance limitado de la búsqueda de hiperparámetros (grillas ampliadas pero no una '
        'búsqueda exhaustiva ni Bayesian optimization): dado el tamaño reducido del dataset de '
        'consenso, una búsqueda más agresiva habría aumentado el riesgo de sobreajustar la '
        'selección de hiperparámetros al ruido de folds pequeños.',
        'BETO (fine-tuning de transformer) omitido en esta iteración: no hay GPU ni `torch` '
        'instalado en el entorno de ejecución disponible; el pipeline ya soporta `--skip-beto` '
        'para este caso y queda pendiente para cuando se disponga de cómputo adecuado.',
        'Serialización de modelos con pickle en vez de skops (formato por defecto de mlflow '
        'reciente): skops rechaza por seguridad tipos como CalibratedClassifierCV/KFold salvo que '
        'se declaren explícitamente confiables; se usó pickle porque el modelo se genera y '
        'consume dentro del mismo equipo, no se distribuye a terceros.',
    ]:
        doc.add_paragraph(texto, style='List Bullet')

    heading(doc, '4.2 Riesgos técnicos identificados y mitigación', nivel=2)
    filas_riesgos = [
        ['Confusión "más datos" vs. "mejor etiqueta" al comparar corpus de tamaños distintos',
         'Se evaluaron ambas fuentes de etiqueta sobre las mismas 1,297 filas y los mismos folds '
         '(diseño pareado), aislando el efecto de la calidad del etiquetado.'],
        ['Clases raras con soporte insuficiente para CV estratificada (QUE=4, SEG=5 registros '
         'totales en el dataset de consenso)',
         'Número de folds adaptativo según la clase minoritaria del train; calibración de '
         'probabilidades condicionada a soporte suficiente; F1-macro con zero_division=0 y '
         'labels fijas para evitar denominadores inconsistentes entre folds.'],
        ['Poder estadístico limitado con n=1,297 y solo 25 folds (parcialmente correlacionados '
         'entre repeticiones de RepeatedKFold)',
         'Se reportan intervalos de confianza y ambas pruebas (paramétrica y no paramétrica) en '
         'vez de un solo p-valor puntual; se recomienda ampliar la anotación en Fase 2 para '
         'reducir esta incertidumbre, especialmente en clases raras.'],
        ['Riesgo de sobreajuste del vocabulario TF-IDF (hasta 15,000 features) sobre un corpus '
         'pequeño (1,037 filas de entrenamiento por fold)',
         'El vectorizador se reajusta dentro de cada fold usando solo el texto de entrenamiento '
         '(sin fuga de información hacia el test); se recomienda monitorear esta métrica al '
         'ampliar el dataset.'],
        ['Incompatibilidad de versiones del entorno de MLflow (>=2.11) con el backend de '
         'archivos local documentado en CLAUDE.md',
         'Documentado el uso de la variable de entorno MLFLOW_ALLOW_FILE_STORE=true; se '
         'recomienda fijar la versión de mlflow en requirements.txt o migrar a un backend '
         'sqlite en una futura iteración.'],
        ['Deriva de documentación (READMEs desactualizados/duplicados, sin bitácora de cambios)',
         'Consolidación de README.md/CLAUDE.md, eliminación de duplicados y creación de '
         'CHANGELOG.md como registro vivo de decisiones y cambios.'],
    ]
    tabla_simple(doc, ['Riesgo identificado', 'Mitigación aplicada'], filas_riesgos)

    RUTA_SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(RUTA_SALIDA))
    print(f"✓ Informe guardado: {RUTA_SALIDA}")


if __name__ == '__main__':
    main()
